from docx import Document

# Read the template
doc = Document('e:/trae/2LIMS/2 检测报告.docx')

print('=== DOCX Template Analysis ===\n')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}\n')

print('=== All Paragraphs ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'{i}: {p.text[:150]}')

print('\n=== Table 0 (Cover/Header) ===')
if len(doc.tables) > 0:
    t0 = doc.tables[0]
    for i, row in enumerate(t0.rows):
        cells_text = ' | '.join([cell.text.strip().replace('\n', ' ')[:80] for cell in row.cells])
        print(f'Row {i}: {cells_text}')

print('\n=== Table 1 (Client Info) ===')
if len(doc.tables) > 1:
    t1 = doc.tables[1]
    for i, row in enumerate(t1.rows):
        cells_text = ' | '.join([cell.text.strip().replace('\n', ' ')[:80] for cell in row.cells])
        print(f'Row {i}: {cells_text}')

print('\n=== Table 2 (Test Results) ===')
if len(doc.tables) > 2:
    t2 = doc.tables[2]
    for i, row in enumerate(t2.rows):
        cells_text = ' | '.join([cell.text.strip().replace('\n', ' ')[:80] for cell in row.cells])
        print(f'Row {i}: {cells_text}')
