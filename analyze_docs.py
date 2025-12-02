import os
from docx import Document
from pypdf import PdfReader

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading DOCX {file_path}: {str(e)}"

def read_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        full_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading PDF {file_path}: {str(e)}"

files_to_analyze = [
    r"e:\trae\2LIMS\56 样品检测委托单-国轻.docx",
    r"e:\trae\2LIMS\ALTC2509034复合材料力学性能.pdf",
    r"e:\trae\2LIMS\ALTC2510007G-奇瑞金相.pdf",
    r"e:\trae\2LIMS\55 检测任务通知单 - 国轻.docx"
]

print("=== START ANALYSIS ===\n")

for file_path in files_to_analyze:
    print(f"--- Analyzing: {os.path.basename(file_path)} ---")
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        continue
        
    if file_path.lower().endswith('.docx'):
        content = read_docx(file_path)
    elif file_path.lower().endswith('.pdf'):
        content = read_pdf(file_path)
    else:
        content = "Unsupported file format"
        
    print(content[:2000]) # Print first 2000 chars to avoid overwhelming output
    print("\n" + "="*50 + "\n")

print("=== END ANALYSIS ===")
